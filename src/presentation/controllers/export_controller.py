from presentation.states.export_state import ExportState, FileFormat
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path
import base64
import io
import os
import tempfile
import pypdfium2 as pdfium
import shutil
import numpy as np
from PIL import Image as IMG
import win32api

from flet import Page, margin as mg, SnackBar, Text, SnackBarBehavior
from presentation.states.active_file_state import ActiveFileState
from presentation.controllers.controller import Controller, Priority
from presentation.states.render_state import RenderState
from presentation.states.sidebar_hide_state import SideBarHideState
from presentation.states.dialogs_state import Dialogs, DialogState
import docx2pdf as d2f

class ExportController(Controller):
    priority = Priority.NONE
    def __init__(self, page: Page):
        self.page = page

        self.export_state = ExportState()
        self.render_state = RenderState()
        self.sbh_state = SideBarHideState()
        self.dia_state = DialogState()
        self.af_state = ActiveFileState()

        self.export_state.on_export = lambda: self.export(to_print=False)
        self.export_state.on_print = lambda: self.export(to_print=True)

        self.documents_dir = (Path.home() / "Documents").absolute()

    def export(self, to_print: bool):
        file_format: FileFormat = self.export_state.format
        margin: bool = self.export_state.margin
        titleblock_enable: bool = self.export_state.titleblock_enable
        proj_name: str = self.export_state.proj_name
        creator: str = self.export_state.creator
        date: str = self.export_state.date

        self.key_name = ""
        for key in self.render_state.image.keys():
            if key == self.af_state.active.title:
                self.key_name = key
                break

        self.dia_state.state = Dialogs.CLOSE

        export_message = "Preparing file for printing. Please wait!" if to_print else "Exporting the document. Please wait!"

        self.page.open(
            SnackBar(
                content=Text(export_message), 
                behavior=SnackBarBehavior.FLOATING, 
                duration=6000,
                show_close_icon=True,
                margin=mg.all(12) if not self.sbh_state.state.value else mg.only(left=212, top=12, right=12, bottom=12),
            )
        )

        output_filename = ""

        match file_format:
            case FileFormat.PDF:
                output_filename = self.export_to_file(margin, titleblock_enable, proj_name, creator, date, 1)
            case FileFormat.PNG:
                output_filename = self.export_to_file(margin, titleblock_enable, proj_name, creator, date, 2)
            case FileFormat.DOCX:
                output_filename = self.export_to_file(margin, titleblock_enable, proj_name, creator, date)
            case FileFormat.RAW_PNG:
                output_filename = self.export_to_png()
        
        if not to_print:
            message = f"Successfully exported to {output_filename}!" if output_filename != "" else f"There was a problem exporting the file!"

            self.page.open(
                SnackBar(
                    content=Text(message), 
                    behavior=SnackBarBehavior.FLOATING, 
                    duration=6000,
                    show_close_icon=True,
                    margin=mg.all(12) if not self.sbh_state.state.value else mg.only(left=212, top=12, right=12, bottom=12),
                    action="Open" if output_filename != "" else None,
                    on_action=lambda e: os.startfile(output_filename) if output_filename != "" else None
                )
            )
        else:
            try:
                win32api.ShellExecute(0, "print", output_filename, None, ".", 0)
            except:
                self.page.open(
                    SnackBar(
                        content=Text("Cannot process the printing of the file. No printers found."), 
                        behavior=SnackBarBehavior.FLOATING, 
                        duration=6000,
                        show_close_icon=True,
                        margin=mg.all(12) if not self.sbh_state.state.value else mg.only(left=212, top=12, right=12, bottom=12),
                        action="Open" if output_filename != "" else None,
                        on_action=lambda e: os.startfile(output_filename) if output_filename != "" else None
                    )
                )
    
    def export_to_file(self, margin: bool, titleblock_enable: bool, proj_name: str, creator: str, date: str, is_pdf = 0):
        if self.key_name == "":
            self.key_name = "New"
        
        image_data = base64.b64decode(self.render_state.image[self.key_name])
        image_stream = io.BytesIO(image_data)

        image_stream = self.center_image(image_stream)
        
        try:
            doc: Document = None
            if margin and titleblock_enable:
                doc = Document("src/assets/full.docx")

                table = doc.tables[0]

                # Insert image in merged cell 0,0 and 0,1
                image_cell = table.cell(0, 0)
                paragraph = image_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(6))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

                # Fill title, creator, and date
                table.cell(1, 0).text = table.cell(1, 0).text.replace('{TITLE}', proj_name)
                table.cell(1, 1).text = table.cell(1, 1).text.replace('{CREATOR}', creator)
                table.cell(2, 1).text = table.cell(2, 1).text.replace('{DATE}', date)
            elif not margin and titleblock_enable:
                doc = Document("src/assets/no_margin.docx")

                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(6))

                # Footer table handling
                section = doc.sections[0]
                footer = section.footer
                table = footer.tables[0]

                table.cell(0, 0).text = table.cell(0, 0).text.replace('{TITLE}', proj_name)
                table.cell(0, 1).text = table.cell(0, 1).text.replace('{CREATOR}', creator)
                table.cell(1, 1).text = table.cell(1, 1).text.replace('{DATE}', date)
            elif margin and not titleblock_enable:
                doc = Document("src/assets/no_titlebar.docx")

                table = doc.tables[0]

                # Insert image into the single large table cell
                image_cell = table.cell(0, 0)
                paragraph = image_cell.paragraphs[0]
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(6))
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            else:
                doc = Document("src/assets/plain.docx")
                paragraph = doc.add_paragraph()
                paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                run = paragraph.add_run()
                run.add_picture(image_stream, width=Inches(6))
            
            if is_pdf == 0:
                doc.save(self.documents_dir / f"{self.key_name}.docx")
                return f"{self.key_name}.docx"

            elif is_pdf == 1:
                with tempfile.TemporaryDirectory() as tempdir:
                    docx_path = os.path.join(tempdir, "temp.docx")
                    pdf_path = self.documents_dir / f"{self.key_name}.pdf"
            
                    doc.save(docx_path)

                    # Convert DOCX to PDF
                    d2f.convert(docx_path, pdf_path)
                
                    return pdf_path
            else:
                output_pdf_path = "copied_temp.pdf"
                with tempfile.TemporaryDirectory() as tempdir:
                    docx_path = os.path.join(tempdir, "temp.docx")
                    pdf_path = os.path.join(tempdir, "temp.pdf")
            
                    doc.save(docx_path)

                    # Convert DOCX to PDF
                    d2f.convert(docx_path, pdf_path)
                    shutil.copy2(pdf_path, output_pdf_path)

                    pdf = pdfium.PdfDocument(output_pdf_path)

                    image = pdf[0].render(scale=4).to_pil()
                    image.save(self.documents_dir / f"{self.key_name}.png", format='PNG')

                return f"{self.key_name}.png"
        except:
            return ""
    
    def export_to_png(self):
        if self.key_name == "":
            self.key_name = "New"

        image_data = base64.b64decode(self.render_state.image[self.key_name])
        image_stream = io.BytesIO(image_data)

        image_stream = self.center_image(image_stream)

        img = IMG.open(image_stream)
    
        img.save(self.documents_dir / f"{self.key_name}.png", format='PNG')

        return f"{self.key_name}.png"
    
    def center_image(self, image_stream: io.BytesIO, bg_color: tuple = None, target_width_inches: float = 12, dpi: int = 96):
        try:
            # Ensure stream is at the start
            image_stream.seek(0)
            
            # Load image
            img = IMG.open(image_stream)
            
            # Convert to RGBA for transparency
            if img.mode != "RGBA":
                img = img.convert("RGBA")
        except Exception as e:
            raise ValueError(f"Error loading image: {e}")
        
        # Step 1: Find the bounding box of the content
        img_array = np.array(img)
        
        if img.mode == "RGBA" and bg_color is None:
            # Use alpha channel for PNGs
            alpha = img_array[:, :, 3]
            non_transparent = np.where(alpha > 0)
            if len(non_transparent[0]) == 0:
                raise ValueError("No non-transparent content found")
            y_min, y_max = np.min(non_transparent[0]), np.max(non_transparent[0])
            x_min, x_max = np.min(non_transparent[1]), np.max(non_transparent[1])
        else:
            # Use background color
            if bg_color is None:
                bg_color = (255, 255, 255)
            non_bg = np.any(img_array[:, :, :3] != bg_color, axis=2)
            coords = np.where(non_bg)
            if len(coords[0]) == 0:
                raise ValueError("No content found (all pixels match background)")
            y_min, y_max = np.min(coords[0]), np.max(coords[0])
            x_min, x_max = np.min(coords[1]), np.max(coords[1])
        
        # Calculate content size
        content_width = x_max - x_min + 1
        content_height = y_max - y_min + 1
        
        # Step 2: Center the content
        img_width, img_height = img.size
        new_x = (img_width - content_width) // 2
        new_y = (img_height - content_height) // 2
        
        centered_img = IMG.new("RGBA", (img_width, img_height), (0, 0, 0, 0))
        content = img.crop((x_min, y_min, x_max + 1, y_max + 1))
        centered_img.paste(content, (new_x, new_y))
        
        # Step 3: Trim transparent pixels
        trimmed_bbox = centered_img.getbbox()  # Returns (x_min, y_min, x_max, y_max) of non-transparent area
        if trimmed_bbox is None:
            raise ValueError("No content after centering")
        trimmed_img = centered_img.crop(trimmed_bbox)
        
        # Get trimmed dimensions
        trimmed_width, trimmed_height = trimmed_img.size
        
        # Step 4: Scale if needed
        target_width_pixels = int(target_width_inches * dpi)
        if trimmed_width < target_width_pixels:
            scale_factor = target_width_pixels / trimmed_width
            new_width = target_width_pixels
            new_height = int(trimmed_height * scale_factor)
            scaled_img = trimmed_img.resize((new_width, new_height), IMG.Resampling.LANCZOS)
        else:
            scaled_img = trimmed_img
            new_width, new_height = trimmed_width, trimmed_height
        
        # Step 5: Create final image (ensure content is centered in canvas)
        final_canvas_width = max(new_width, img_width)  # Preserve original size if larger
        final_canvas_height = max(new_height, img_height)
        final_img = IMG.new("RGBA", (final_canvas_width, final_canvas_height), (0, 0, 0, 0))
        final_x = (final_canvas_width - new_width) // 2
        final_y = (final_canvas_height - new_height) // 2
        final_img.paste(scaled_img, (final_x, final_y))
        
        # Step 6: Save to BytesIO stream
        output_stream = io.BytesIO()
        final_img.save(output_stream, format="PNG")
        output_stream.seek(0)
        
        return output_stream